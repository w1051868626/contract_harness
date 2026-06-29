"""知识库模块，支持文档管理、智能分块与语义检索。"""

from __future__ import annotations

import json
import re
import tempfile
import uuid
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

# ---- 法律文本逐行解析：中文数字 & 正则（模块级，供 _chunk_law_text 使用） ----
from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from harness.agent.llm import LLMClient, LLMResponse
from harness.core.config import HarnessConfig, LLMConfig
from harness.core.exceptions import ChunkingError
from harness.rag.constants import (
    _ART_PAT_RE,
    _ART_SEARCH_RE,
    _ARTICLE_LAW_RE,
    _CHAPTER_DIVISION_RE,
    _CHAPTER_LAW_RE,
    _CN_DIGIT,
    _DECIMAL_RANGE_RE,
    _DECIMAL_RE,
    _MD_HEADING_RE,
    _MD_SPLIT_RE,
    _NUM_LIST_RE,
    _PAREN_NUM_RE,
    _SECTION_HEAD_RE,
    _SECTION_LAW_RE,
    CHUNK_MAX_CHARS,
    CHUNK_PROMPT,
    CHUNK_SYSTEM_PROMPT,
    EXPANSION_SYSTEM_PROMPT,
    QUERY_EXPANSION_PROMPT,
    DocType,
    MetaKey,
)
from harness.rag.docling_parser import DoclingParser
from harness.rag.embedding import EmbeddingProvider, create_embedding_provider
from harness.rag.reranker import Reranker, create_reranker
from harness.rag.vector_store import Chunk, Document, VectorStore, create_vector_store
from harness.utils.io import normalize_text as _util_normalize
from harness.utils.log import logger

# ---- 法律文本逐行解析辅助（模块级，供 _chunk_law_text 使用） ----

_CN_NUM: dict[str, int] = {
    "零": 0,
    "〇": 0,
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
    "百": 100,
    "千": 1000,
}


def _chinese_to_int(text: str) -> int:
    """中文数字转整数。"""
    if text.isdigit():
        return int(text)
    total = 0
    num = 0
    for c in reversed(text):
        if c == "十":
            if num == 0:
                num = 1
            total += num * 10
            num = 0
        elif c == "百":
            if num == 0:
                num = 1
            total += num * 100
            num = 0
        elif c == "千":
            if num == 0:
                num = 1
            total += num * 1000
            num = 0
        else:
            num = _CN_NUM[c]
    total += num
    return total


_TITLE_RE = re.compile(r"^中华人民共和国.*")
_DATE_RE = re.compile(r"（(\d{4})年(\d+)月(\d+)日.*?通过")


@dataclass
class _LawContext:
    """法律文本逐行解析的上下文（不变字段），供 _flush_law_article 使用。"""

    doc_id: str
    chunk_size: int
    law_name: str | None = None
    pub_date: str | None = None
    chapter: str | None = None
    section: str | None = None
    chunks: list[Chunk] = field(default_factory=list)


@dataclass
class _MdAccumulator:
    """Markdown 分块累加器，封装 buffer/buf_meta/buf_len/idx。"""

    buffer: list[str]
    buf_meta: list[dict[MetaKey, str]]
    buf_len: int
    idx: int
    chunk_size: int
    overlap: int
    doc_id: str
    chunks: list[Chunk]

    def add(self, sec: str, this_meta: dict[MetaKey, str]) -> None:
        sl = len(sec)
        if self.buf_len + sl <= self.chunk_size:
            self.buffer.append(sec)
            self.buf_meta.append(this_meta)
            self.buf_len += sl
        else:
            self.flush()
            if sl <= self.chunk_size:
                self.buffer, self.buf_meta, self.buf_len = [sec], [this_meta], sl
            else:
                for sub in KnowledgeBase._split_segments(sec):
                    if self.buf_len + len(sub) <= self.chunk_size:
                        self.buffer.append(sub)
                        self.buf_meta.append(this_meta)
                        self.buf_len += len(sub)
                    else:
                        self.flush()
                        self.buffer, self.buf_meta, self.buf_len = [sub], [this_meta], len(sub)

    def flush(self) -> None:
        if not self.buffer:
            return
        meta: dict[str, Any] = {}
        for bm in self.buf_meta:
            for key in (MetaKey.CHAPTER, MetaKey.SECTION):
                if bm.get(key):
                    meta.setdefault(key, bm[key])
        arts = KnowledgeBase._md_article_range(self.buffer)
        if arts:
            meta[MetaKey.ARTICLES] = arts
        self.chunks.append(
            KnowledgeBase._make_chunk(self.buffer, self.doc_id, self.idx, metadata=meta)
        )
        self.idx += 1
        carry = KnowledgeBase._carry_overlap(self.buffer, self.overlap)
        self.buffer = carry
        self.buf_meta = self.buf_meta[-len(carry) :] if carry else []
        self.buf_len = sum(len(s) for s in carry)


class KnowledgeBase:
    """知识库，管理文档的添加、分块、嵌入与检索。"""

    def __init__(
        self,
        store: VectorStore,
        embedding: EmbeddingProvider,
        llm: LLMClient | None = None,
        reranker: Reranker | None = None,
        chunk_model: str = "gpt-4o-mini",
        chunk_llm: LLMClient | None = None,
        expansion_llm: LLMClient | None = None,
    ):
        """初始化知识库。"""
        self._store = store
        self._embedding = embedding
        self._llm = llm
        self._reranker = reranker
        self._chunk_model = chunk_model
        self._chunk_llm = chunk_llm
        self._expansion_llm = expansion_llm or llm

    @property
    def store(self) -> VectorStore:
        """返回底层向量存储。"""
        return self._store

    @classmethod
    def from_config(cls, config: HarnessConfig | None = None) -> KnowledgeBase:
        """从 HarnessConfig 创建知识库实例。"""
        cfg = config or HarnessConfig()
        if cfg.use_docling:
            cls.enable_docling()
        store = create_vector_store(cfg.kb_dir)
        embedding = create_embedding_provider(
            provider=cfg.embedding.provider,
            api_key=cfg.embedding.api_key,
            api_base=cfg.embedding.api_base,
            model=cfg.embedding.model,
            proxy=cfg.embedding.proxy,
        )
        reranker = create_reranker(
            provider=cfg.embedding.rerank_provider,
            api_key=cfg.embedding.rerank_api_key,
            api_base=cfg.embedding.rerank_api_base,
            model=cfg.embedding.rerank_model,
        )
        llm = LLMClient(cfg.llm) if cfg.llm.api_key else None
        chunk_cfg = LLMConfig(
            api_key=cfg.llm.chunk_api_key,
            api_base=cfg.llm.chunk_api_base,
            model=cfg.llm.chunk_model,
            proxy=cfg.llm.proxy,
        )
        chunk_llm = LLMClient(chunk_cfg) if chunk_cfg.api_key and chunk_cfg.model else None
        expansion_cfg = LLMConfig(
            api_key=cfg.llm.expansion_api_key,
            api_base=cfg.llm.expansion_api_base,
            model=cfg.llm.expansion_model,
            proxy=cfg.llm.proxy,
        )
        expansion_llm = (
            LLMClient(expansion_cfg) if expansion_cfg.api_key and expansion_cfg.model else None
        )
        logger.debug(
            "KnowledgeBase.from_config: kb_dir={}, embedding={}, reranker={},"
            " chunk_llm={}, expansion_llm={}",
            cfg.kb_dir,
            cfg.embedding.provider,
            cfg.embedding.rerank_provider or "none",
            chunk_cfg.model if chunk_llm else "none",
            expansion_cfg.model if expansion_llm else "none",
        )
        return cls(
            store=store,
            embedding=embedding,
            reranker=reranker,
            llm=llm,
            chunk_llm=chunk_llm,
            chunk_model=cfg.llm.chunk_model,
            expansion_llm=expansion_llm,
        )

    @staticmethod
    def _new_id() -> str:
        """生成短 ID。"""
        return uuid.uuid4().hex[:12]

    @staticmethod
    def _normalize_text(text: str) -> str:
        """清洗文本中的非常规字符。"""
        return _util_normalize(text)

    def add_text(
        self,
        title: str,
        content: str,
        source: str = "",
        metadata: dict[str, Any] | None = None,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        use_ai_chunking: bool = True,
    ) -> str:
        """将文本添加到知识库。"""
        content = KnowledgeBase._normalize_text(content)
        doc_id = KnowledgeBase._new_id()
        doc = Document(
            id=doc_id,
            title=title,
            source=source,
            metadata=metadata or {},
        )
        self._store.add_document(doc)
        chunks = self._resolve_chunks(content, doc_id, chunk_size, chunk_overlap, use_ai_chunking)
        if chunks:
            texts = [c.content for c in chunks]
            embeddings = self._embedding.embed_batch(texts)
            for chunk, emb in zip(chunks, embeddings):
                chunk.embedding = emb
            self._store.add_chunks(chunks)
            logger.debug("添加文本: title={}, doc_id={}, chunks={}", title, doc_id, len(chunks))
        else:
            logger.warning("添加文本无分块: title={}, doc_id={}", title, doc_id)
        return doc_id

    def _resolve_chunks(
        self,
        content: str,
        doc_id: str,
        chunk_size: int,
        chunk_overlap: int,
        use_ai: bool,
    ) -> list[Chunk]:
        """根据配置选择 AI 分块、法律条文分块或传统分块。"""
        if use_ai and self._chunk_llm is not None:
            logger.debug("使用 AI 分块")
            try:
                return self._chunk_with_ai(content, doc_id)
            except (ValueError, json.JSONDecodeError, RuntimeError):
                logger.warning("AI 分块失败，回退到传统分块", exc_info=True)
        try:
            md = self._chunk_markdown(content, doc_id, chunk_size, chunk_overlap)
            if md is not None:
                logger.debug("使用 Markdown 结构化分块: chunks={}", len(md))
                return md
            law = self._chunk_law_text(content, doc_id, chunk_size)
            if law is not None:
                logger.debug("使用逐条法律分块: chunks={}", len(law))
                return law
            legal = self._chunk_legal_text(content, doc_id, chunk_size, chunk_overlap)
            if legal is not None:
                logger.debug("使用法律条文分块: chunks={}", len(legal))
                return legal
            logger.debug("使用通用文本分块")
            return self._chunk_text(content, doc_id, chunk_size, chunk_overlap)
        except Exception as exc:
            # 将各分块策略的异常统一为 ChunkingError
            raise ChunkingError(f"所有分块策略均失败: {exc}") from exc

    def _chunk_with_ai(self, text: str, doc_id: str) -> list[Chunk]:
        """使用 LLM 对文本进行智能分块，超长文本自动分段后合并。"""
        if self._chunk_llm is None:
            raise RuntimeError("chunk_llm 未初始化")
        all_chunks: list[Chunk] = []
        for i in range(0, len(text), CHUNK_MAX_CHARS):
            segment = text[i : i + CHUNK_MAX_CHARS]
            prompt = CHUNK_PROMPT.format(text=segment)
            resp: LLMResponse = self._chunk_llm.chat(
                [
                    {"role": "system", "content": CHUNK_SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                model=self._chunk_model,
                temperature=0.0,
            )
            raw = resp.content.strip()
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
            data = json.loads(raw)
            if not isinstance(data, list):
                raise ValueError("AI chunking did not return a list")
            for item in data:
                content = item.get("content", "").strip()
                if content:
                    all_chunks.append(
                        Chunk(
                            id=KnowledgeBase._new_id(),
                            document_id=doc_id,
                            content=content,
                            chunk_index=len(all_chunks),
                        )
                    )
        logger.debug("AI 分块完成: chunks={}", len(all_chunks))
        return all_chunks

    def add_file(
        self,
        file_path: str | Path,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        use_ai_chunking: bool = True,
    ) -> str:
        """添加文件到知识库。"""
        path = Path(file_path)
        logger.info("添加文件: {}", path.name)
        if path.suffix.lower() == ".zip":
            doc_ids = self.add_zip(path, chunk_size, chunk_overlap, use_ai_chunking)
            return doc_ids[0] if doc_ids else ""
        content = self._parse_file(path)
        return self.add_text(
            title=path.stem,
            content=content,
            source=str(path),
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            use_ai_chunking=use_ai_chunking,
        )

    @staticmethod
    def _extract_zip_texts(path: Path) -> list[tuple[str, str]]:
        """解压 zip，返回 (内部文件名, 解析文本) 列表。"""
        supported = {".txt", ".md", ".json", ".pdf", ".docx"}
        results: list[tuple[str, str]] = []
        with zipfile.ZipFile(path) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                ext = Path(info.filename).suffix.lower()
                if ext not in supported:
                    continue
                try:
                    raw = zf.read(info.filename)
                    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                        tmp.write(raw)
                        tmp_path = Path(tmp.name)
                    content = KnowledgeBase._parse_file(tmp_path)
                    tmp_path.unlink(missing_ok=True)
                    if content.strip():
                        results.append((info.filename, content))
                except (json.JSONDecodeError, KeyError, OSError, zipfile.BadZipFile):
                    logger.warning("ZIP 中文件解析失败: {}", info.filename, exc_info=True)
                    continue
        return results

    def add_zip(
        self,
        path: Path,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        use_ai_chunking: bool = True,
    ) -> list[str]:
        """解压 zip 并以内部文件名为标题分别导入。"""
        doc_ids: list[str] = []
        logger.debug("解压 zip 导入: path={}", path)
        for filename, content in KnowledgeBase._extract_zip_texts(path):
            try:
                title = Path(filename).stem
                doc_id = self.add_text(
                    title=title,
                    content=content,
                    source=str(path),
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    use_ai_chunking=use_ai_chunking,
                )
                doc_ids.append(doc_id)
            except (json.JSONDecodeError, KeyError, OSError, zipfile.BadZipFile):
                logger.warning("ZIP 中文件处理失败: filename={}", filename, exc_info=True)
                continue
        return doc_ids

    @staticmethod
    def _parse_txt(path: Path) -> str:
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _parse_json(path: Path) -> str:
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, list):
            return "\n".join(json.dumps(item, ensure_ascii=False) for item in data)
        if isinstance(data, dict):
            return json.dumps(data, ensure_ascii=False, indent=2)
        return str(data)

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except (PdfReadError, KeyError, IndexError):
            logger.warning("PDF 解析失败，按文本读取: {}", path)
            return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _parse_docx(path: Path) -> str:
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    _PARSERS: dict[str, Callable[[Path], str]] = {
        ".txt": _parse_txt,
        ".md": _parse_txt,
        ".json": _parse_json,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
    }

    @staticmethod
    def _parse_file(path: Path) -> str:
        """解析文件内容（支持 txt/md/json/pdf/docx，可选 Docling）。"""
        suffix = path.suffix.lower()
        logger.debug("解析文件: path={}, suffix={}", path.name, suffix)

        docling_parser = getattr(KnowledgeBase, "_docling_parser", None)
        if docling_parser and docling_parser.available and docling_parser.supports(path):
            logger.info("Docling 解析: path={}, suffix={}", path.name, suffix)
            try:
                md = docling_parser.parse_to_markdown(path)
                if md.strip():
                    logger.info("Docling 解析成功: path={} ({} 字符)", path.name, len(md))
                    return md
                logger.warning("Docling 返回空内容，回退 path={}", path.name)
            except RuntimeError as e:
                logger.warning("Docling 解析失败，回退 path={}: {}", path.name, e)

        handler = KnowledgeBase._PARSERS.get(suffix)
        if handler is not None:
            return handler(path)
        return path.read_text(encoding="utf-8")

    @classmethod
    def enable_docling(cls) -> None:
        """启用 Docling 解析器（在 KnowledgeBase 级共享单例）。"""
        parser = DoclingParser()
        if parser.available:
            cls._docling_parser = parser
            logger.info("Docling 解析器已启用")
        else:
            logger.warning("Docling 解析器不可用（未安装或初始化失败）")
            cls._docling_parser = None

    def query(self, text: str, top_k: int = 5, expansion_threshold: float = 0.6) -> list[Chunk]:
        """语义检索最相关的文本块（支持可选的 rerank 精排和 AI 扩展检索词）。

        Args:
            text: 查询文本。
            top_k: 返回结果数量。
            expansion_threshold: AI 扩展检索词的相似度阈值。
                当最高分低于此值时，用 LLM 生成同义查询重试。
                设为 0 可禁用扩展。
        """
        logger.debug("检索: query={}, top_k={}", text[:50], top_k)
        candidates = self._search_single(text, top_k)
        if not candidates:
            return []

        max_score = candidates[0].score
        if (
            expansion_threshold > 0
            and max_score < expansion_threshold
            and self._expansion_llm is not None
        ):
            logger.info(
                "检索结果分数偏低 (max={:.3f} < {:.2f})，使用 AI 扩展检索词",
                max_score,
                expansion_threshold,
            )
            expanded = self._expand_query(text)
            if len(expanded) > 1:
                for q in expanded[1:]:
                    extra = self._search_single(q, top_k)
                    candidates = self._merge_results(candidates, extra, top_k)
                logger.debug("扩展检索完成: results={}", len(candidates))
        return candidates[:top_k]

    def _search_single(self, text: str, top_k: int) -> list[Chunk]:
        """单次语义检索（含可选的 rerank）。"""
        query_emb = self._embedding.embed(text)
        candidates = self._store.search(query_emb, top_k=top_k * 2 if self._reranker else top_k)
        if self._reranker and len(candidates) > 1:
            return self._reranker.rerank(text, candidates, top_k=top_k)
        return candidates[:top_k]

    def _expand_query(self, query: str, num_variants: int = 2) -> list[str]:
        """使用 LLM 生成同义检索词变体，返回包含原始查询的列表。"""
        if self._expansion_llm is None:
            msg = "_expand_query 需要 LLMClient，但未配置"
            raise RuntimeError(msg)
        prompt = QUERY_EXPANSION_PROMPT.format(query=query, num_variants=num_variants)
        try:
            resp: LLMResponse = self._expansion_llm.chat(
                [
                    {"role": "system", "content": EXPANSION_SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
            )
            variants = [line.strip() for line in resp.content.strip().split("\n") if line.strip()]
            result = [query] + variants[:num_variants]
            logger.debug("扩展检索词: {}", result)
            return result
        except Exception:
            # LLM 调用可能抛多种 openai 异常，全部降级为原始查询
            logger.warning("AI 扩展检索词失败，使用原始查询", exc_info=True)
            return [query]

    @staticmethod
    def _merge_results(
        existing: list[Chunk],
        new_results: list[Chunk],
        top_k: int,
    ) -> list[Chunk]:
        """合并两组检索结果，按 id 去重并保留最高分。"""
        seen: dict[str, Chunk] = {}
        for chunk in existing + new_results:
            if chunk.id in seen:
                seen[chunk.id].score = max(seen[chunk.id].score, chunk.score)
            else:
                seen[chunk.id] = chunk
        merged = sorted(seen.values(), key=lambda c: c.score, reverse=True)
        return merged[:top_k]

    def clear_cache(self) -> None:
        """清除知识库内部缓存（当前无缓存，预留接口）。"""

    def list_documents(self) -> list[Document]:
        """列出所有文档。"""
        docs = self._store.list_documents()
        logger.debug("列出文档: count={}", len(docs))
        return docs

    def delete_document(self, document_id: str) -> None:
        """删除指定文档及其分块。"""
        logger.debug("删除文档: doc_id={}", document_id)
        self._store.delete_document(document_id)

    @staticmethod
    def _detect_md_heading_meta(first_line: str) -> dict[MetaKey, str]:
        """从 Markdown 标题行提取章节元数据。"""
        m: dict[MetaKey, str] = {}
        raw = first_line.lstrip("#").strip()
        if _CHAPTER_DIVISION_RE.match(raw):
            m[MetaKey.CHAPTER] = raw
        elif _NUM_LIST_RE.match(raw):
            m[MetaKey.CHAPTER] = raw
        elif _DECIMAL_RE.match(raw):
            m[MetaKey.CHAPTER] = raw
        if _SECTION_HEAD_RE.match(raw):
            m[MetaKey.SECTION] = raw
        elif _PAREN_NUM_RE.match(raw):
            m[MetaKey.SECTION] = raw
        elif _DECIMAL_RANGE_RE.match(raw):
            m[MetaKey.SECTION] = raw
        return m

    @staticmethod
    def _chunk_markdown(
        text: str,
        doc_id: str,
        chunk_size: int,
        overlap: int,
    ) -> list[Chunk] | None:
        """以 # 标题为界的结构化分块，无 # 则返回 None。"""
        if not _MD_HEADING_RE.search(text):
            return None

        sections = [s.strip() for s in _MD_SPLIT_RE.split(text.strip()) if s.strip()]
        acc = _MdAccumulator(
            buffer=[],
            buf_meta=[],
            buf_len=0,
            idx=0,
            chunk_size=chunk_size,
            overlap=overlap,
            doc_id=doc_id,
            chunks=[],
        )

        for sec in sections:
            first_line = sec.split("\n")[0].strip()
            this_meta = KnowledgeBase._detect_md_heading_meta(first_line)
            if not this_meta and acc.buf_meta:
                this_meta = dict(acc.buf_meta[-1])
            elif acc.buf_meta and MetaKey.CHAPTER in acc.buf_meta[-1]:
                this_meta.setdefault(MetaKey.CHAPTER, acc.buf_meta[-1][MetaKey.CHAPTER])

            ch = this_meta.get(MetaKey.CHAPTER)
            if ch and acc.buf_len > 0 and ch != acc.buf_meta[0].get(MetaKey.CHAPTER):
                acc.flush()
                acc.buffer, acc.buf_meta, acc.buf_len = [sec], [this_meta], len(sec)
                continue
            sc = this_meta.get(MetaKey.SECTION)
            if sc and acc.buf_len > 0 and sc != acc.buf_meta[-1].get(MetaKey.SECTION):
                acc.flush()
                acc.buffer, acc.buf_meta, acc.buf_len = [sec], [this_meta], len(sec)
                continue

            acc.add(sec, this_meta)

        acc.flush()
        return acc.chunks

    @staticmethod
    def _split_keep_separator(text: str, pattern: str) -> list[str]:
        """按正则分割，分隔符保留在后续片段开头。"""
        parts = re.split(f"({pattern})", text)
        if len(parts) < 2:
            return [text]
        result: list[str] = []
        buf = parts[0]
        i = 1
        while i < len(parts):
            buf += parts[i]
            if i + 1 < len(parts):
                buf += parts[i + 1]
            result.append(buf)
            buf = ""
            i += 2
        if buf.strip():
            result.append(buf)
        return result

    @staticmethod
    def _hierarchical_split(text: str, separators: list[str], chunk_size: int) -> list[str]:
        """递归层级文本分割器（模拟 RecursiveCharacterTextSplitter）。

        从最高优先级分隔符开始尝试；若分割后片段仍超 chunk_size，
        用下一级分隔符递归处理。片段不合并——每个片段即为一个 chunk。
        """
        if not text.strip():
            return []
        if len(text) <= chunk_size or not separators:
            return [text]

        sep = separators[0]
        rest = separators[1:]

        if sep == r"\n":
            parts = [p for p in text.split("\n") if p.strip()]
            if not parts:
                return [text]
        else:
            parts = KnowledgeBase._split_keep_separator(text, sep)

        result: list[str] = []
        for part in parts:
            if len(part) > chunk_size and rest:
                result.extend(KnowledgeBase._hierarchical_split(part, rest, chunk_size))
            else:
                result.append(part)
        return [p for p in result if p.strip()]

    @staticmethod
    def _extract_law_metadata(text: str, filename: str = "") -> dict[str, Any]:
        """从法律文本中提取元数据（法律名称、生效日期等）。"""
        meta: dict[str, Any] = {
            MetaKey.DOC_TYPE: DocType.LAW,
            MetaKey.SOURCE_FILE: filename,
            MetaKey.LAW_NAME: "",
            MetaKey.EFFECTIVE_DATE: "",
        }

        name_pats = [
            r"(中华人民共和国\S+?(?:法典|法|条例|规定|办法|决定|解释))",
            r"(《[^》]+》)",
            r"^(.+(?:法典|法|条例|规定|办法|决定|解释))",
        ]
        for pat in name_pats:
            m = re.search(pat, text[:500])
            if m:
                meta[MetaKey.LAW_NAME] = m.group(1).strip("《》")
                break

        if not meta[MetaKey.LAW_NAME] and filename:
            meta[MetaKey.LAW_NAME] = re.sub(r"\.\w+$", "", filename)

        date_pats = [
            r"(\d{4}年\d{1,2}月\d{1,2}日)(?:起)?施行",
            r"自(\d{4}年\d{1,2}月\d{1,2}日)起施行",
            r"(\d{4}-\d{2}-\d{2})",
        ]
        for pat in date_pats:
            m = re.search(pat, text[:2000])
            if m:
                meta[MetaKey.EFFECTIVE_DATE] = m.group(1)
                break

        return meta

    @staticmethod
    def _extract_case_metadata(text: str, filename: str = "") -> dict[str, Any]:
        """从指导案例文本中提取元数据。"""
        meta: dict[str, Any] = {
            MetaKey.DOC_TYPE: DocType.CASE,
            MetaKey.SOURCE_FILE: filename,
            MetaKey.CASE_NUMBER: "",
            MetaKey.CASE_TITLE: "",
            MetaKey.GUIDING_NUMBER: "",
            MetaKey.KEYWORDS: "",
        }

        g = re.search(r"指导案例(\d+)号", text[:500])
        if g:
            meta[MetaKey.GUIDING_NUMBER] = f"指导案例{g.group(1)}号"

        c = re.search(r"[（(]\d{4}[）)][^\s]*\d+号", text[:1000])
        if c:
            meta[MetaKey.CASE_NUMBER] = c.group(0)

        kw = re.search(r"关键词[：:]\s*(.+?)(?:\n|$)", text[:1000])
        if kw:
            keywords = re.split(r"[；;,，\s]+", kw.group(1).strip())
            kw_list = [k.strip() for k in keywords if k.strip()]
            meta[MetaKey.KEYWORDS] = "；".join(kw_list)

        if not meta[MetaKey.CASE_TITLE] and filename:
            meta[MetaKey.CASE_TITLE] = re.sub(r"\.\w+$", "", filename)

        return meta

    @staticmethod
    def _inject_contextual_header(
        content: str,
        metadata: dict[str, Any],
        doc_meta: dict[str, Any],
    ) -> str:
        """在 chunk 内容前插入结构化上下文标头，提升 Embedding 质量。

        标头格式: [法律名称: X | 章节: Y | 条号: 第Z条]
        """
        parts: list[str] = []

        law_name = doc_meta.get(MetaKey.LAW_NAME, "") or metadata.get(MetaKey.LAW_NAME, "")
        if law_name:
            parts.append(f"法律名称: {law_name}")

        chapter = metadata.get(MetaKey.CHAPTER, "")
        if chapter:
            parts.append(f"章节: {chapter}")

        article = metadata.get(MetaKey.ARTICLES, "")
        if article:
            parts.append(f"条号: {article}")

        guiding = doc_meta.get(MetaKey.GUIDING_NUMBER, "")
        if guiding:
            parts.append(f"案例: {guiding}")

        if parts:
            header = "[" + " | ".join(parts) + "]\n"
            return header + content

        return content

    @staticmethod
    def _chunk_law_text(
        text: str,
        doc_id: str,
        chunk_size: int,
    ) -> list[Chunk] | None:
        """法律文本逐行解析分块：按条聚合，超长条递归切分。

        逐行扫描，识别法律名称、发布日期、章、节、条号；
        每条内容聚合成一个 buffer，超 chunk_size 时按分隔符
        （空行→换行→句号→分号→逗号）递归切分。
        非法律文本返回 None。
        """
        if not _ART_SEARCH_RE.search(text):
            return None

        lines = text.splitlines()
        ctx = _LawContext(doc_id=doc_id, chunk_size=chunk_size)
        cur_article: str | None = None
        cur_article_no: int | None = None
        buffer: list[str] = []
        idx = 0

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 法律名称
            if ctx.law_name is None and _TITLE_RE.match(line):
                ctx.law_name = line
                continue

            # 发布日期
            if ctx.pub_date is None:
                m = _DATE_RE.search(line)
                if m:
                    y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
                    ctx.pub_date = f"{y}-{mo:02d}-{d:02d}"

            # 章
            m = _CHAPTER_LAW_RE.match(line)
            if m:
                ctx.chapter = m.group(1)
                continue

            # 节
            m = _SECTION_LAW_RE.match(line)
            if m:
                ctx.section = m.group(1)
                continue

            # 条
            m = _ARTICLE_LAW_RE.match(line)
            if m:
                idx = KnowledgeBase._flush_law_article(
                    buffer,
                    cur_article,
                    cur_article_no,
                    ctx,
                    idx,
                )
                cur_article = m.group(1)
                cur_article_no = _chinese_to_int(m.group(2))
                buffer = [line]
            else:
                buffer.append(line)

        idx = KnowledgeBase._flush_law_article(
            buffer,
            cur_article,
            cur_article_no,
            ctx,
            idx,
        )

        return ctx.chunks

    @staticmethod
    def _split_recursive(text: str, chunk_size: int) -> list[str]:
        """递归分割长文本，按空行→换行→句号→分号→逗号 优先级切分。"""
        if len(text) <= chunk_size:
            return [text]
        separators = ["\n\n", "\n", "。", "；", "，"]
        for sep in separators:
            parts = text.split(sep)
            if len(parts) > 1:
                result: list[str] = []
                buf = ""
                for p in parts:
                    if not p.strip():
                        if buf:
                            result.append(buf)
                            buf = ""
                        continue
                    candidate = (buf + sep + p) if buf else p
                    if len(candidate) <= chunk_size:
                        buf = candidate
                    else:
                        if buf:
                            result.append(buf)
                        if len(p) <= chunk_size:
                            buf = p
                        else:
                            sub = KnowledgeBase._split_recursive(p, chunk_size)
                            result.extend(sub)
                            buf = ""
                if buf:
                    result.append(buf)
                return result if result else [text]
        return [text]

    @staticmethod
    def _md_article_range(segments: list[str]) -> str:
        """从 Markdown 段列表中提取条文范围。"""
        arts = []
        for s in segments:
            for m in _ART_PAT_RE.finditer(s):
                arts.append(int(m.group(1)) if m.group(1).isdigit() else m.group(1))
        if not arts:
            return ""
        return f"第{arts[0]}条" if len(arts) == 1 else f"第{arts[0]}条—第{arts[-1]}条"

    @staticmethod
    def _flush_law_article(
        buffer: list[str],
        cur_article: str | None,
        cur_article_no: int | None,
        ctx: _LawContext,
        idx: int,
    ) -> int:
        """将法律文本当前条 flush 为一个或多个 Chunk。"""
        if not cur_article:
            return idx
        content = "\n".join(buffer)
        pieces = KnowledgeBase._split_recursive(content, ctx.chunk_size)
        meta_base: dict[str, Any] = {}
        if ctx.law_name:
            meta_base[MetaKey.LAW_NAME] = ctx.law_name
        if ctx.chapter:
            meta_base[MetaKey.CHAPTER] = ctx.chapter
        if ctx.section:
            meta_base[MetaKey.SECTION] = ctx.section
        meta_base[MetaKey.ARTICLES] = cur_article
        if cur_article_no is not None:
            meta_base[MetaKey.ARTICLE_NO] = cur_article_no
        if ctx.pub_date:
            meta_base[MetaKey.EFFECTIVE_DATE] = ctx.pub_date
        meta_base[MetaKey.CHUNK_TOTAL] = len(pieces)
        for piece in pieces:
            ctx.chunks.append(
                Chunk(
                    id=KnowledgeBase._new_id(),
                    document_id=ctx.doc_id,
                    content=KnowledgeBase._align_chunk_end(piece),
                    chunk_index=idx,
                    metadata=dict(meta_base),
                )
            )
            idx += 1
        return idx

    @staticmethod
    def _chunk_legal_text(
        text: str,
        doc_id: str,
        chunk_size: int,
        overlap: int,
    ) -> list[Chunk] | None:
        """法律条文层级回退分块（编→章→节→条→款→项→换行）。

        作为 _chunk_law_text 的 fallback，处理非规范格式的法律文本。
        """
        if not _ART_SEARCH_RE.search(text):
            return None

        separators = [
            rf"\n第{_CN_DIGIT}+编",
            rf"\n第{_CN_DIGIT}+章",
            rf"\n第{_CN_DIGIT}+节",
            rf"\n第{_CN_DIGIT}+条",
            rf"\n{_CN_DIGIT}+[、．]",
            rf"\n[（(]{_CN_DIGIT}+[）)]",
            r"\n",
        ]

        parts = KnowledgeBase._hierarchical_split(text.strip(), separators, chunk_size)
        if not parts:
            return None

        chunks: list[Chunk] = []
        idx = 0
        cur_chapter = ""
        cur_section = ""
        prev_tail = ""

        for part in parts:
            content = prev_tail + part if prev_tail else part
            meta: dict[str, Any] = {}

            first_ch = ""
            first_sec = ""
            for line in part.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if _CHAPTER_DIVISION_RE.match(line) and not first_ch:
                    first_ch = line
                if _SECTION_HEAD_RE.match(line) and not first_sec:
                    first_sec = line

            if first_ch:
                cur_chapter = first_ch
                cur_section = ""
            if first_sec:
                cur_section = first_sec

            if cur_chapter:
                meta[MetaKey.CHAPTER] = cur_chapter
            if cur_section:
                meta[MetaKey.SECTION] = cur_section

            arts = _ART_PAT_RE.findall(part)
            if arts:
                meta[MetaKey.ARTICLES] = (
                    f"第{arts[0]}条" if len(arts) == 1 else f"第{arts[0]}条—第{arts[-1]}条"
                )

            chunks.append(
                Chunk(
                    id=KnowledgeBase._new_id(),
                    document_id=doc_id,
                    content=KnowledgeBase._align_chunk_end(content),
                    chunk_index=idx,
                    metadata=meta,
                )
            )
            idx += 1
            prev_tail = part[-overlap:] if overlap > 0 else ""

        return chunks

    @staticmethod
    def _chunk_text(
        text: str,
        doc_id: str,
        chunk_size: int,
        overlap: int,
    ) -> list[Chunk]:
        """基于段落滑动窗口的传统分块算法。"""
        if not text.strip():
            return []

        segments = KnowledgeBase._split_segments(text)

        if len(segments) == 1 and len(text) <= chunk_size:
            return [
                Chunk(
                    id=KnowledgeBase._new_id(),
                    document_id=doc_id,
                    content=KnowledgeBase._align_chunk_end(text),
                    chunk_index=0,
                )
            ]

        chunks: list[Chunk] = []
        idx = 0
        buffer: list[str] = []
        buffer_len = 0

        for seg in segments:
            seg_len = len(seg)
            if buffer_len + seg_len <= chunk_size:
                buffer.append(seg)
                buffer_len += seg_len
            else:
                if buffer:
                    chunks.append(KnowledgeBase._make_chunk(buffer, doc_id, idx))
                    idx += 1
                    carry = KnowledgeBase._carry_overlap(buffer, overlap)
                    buffer = carry
                    buffer_len = sum(len(s) for s in carry)

                if seg_len > chunk_size:
                    for sub in KnowledgeBase._split_long(seg, chunk_size):
                        chunks.append(
                            Chunk(
                                id=KnowledgeBase._new_id(),
                                document_id=doc_id,
                                content=KnowledgeBase._align_chunk_end(sub),
                                chunk_index=idx,
                            )
                        )
                        idx += 1
                    buffer = []
                    buffer_len = 0
                else:
                    buffer = [seg]
                    buffer_len = seg_len

        if buffer:
            chunks.append(KnowledgeBase._make_chunk(buffer, doc_id, idx))

        return chunks

    @staticmethod
    def _split_segments(text: str) -> list[str]:
        """按空行将文本分割为段落。"""
        raw = re.split(r"\n\s*\n", text.strip())
        return [s.strip() for s in raw if s.strip()]

    @staticmethod
    def _align_chunk_end(content: str) -> str:
        """将内容结尾对齐到句末标点。"""
        matches = list(re.finditer(r"[。！？；.!?;]", content))
        if matches:
            return content[: matches[-1].end()].rstrip()
        return content

    @staticmethod
    def _align_overlap_start(tail: str) -> str:
        """将 overlap 尾部对齐到句子开头。"""
        m = re.search(r"[。！？；.!?;]\s*", tail)
        if m:
            return tail[m.end() :]
        return tail

    @staticmethod
    def _split_long(text: str, chunk_size: int) -> list[str]:
        """将长文本按句子切分。"""
        sentences = re.split(r"(?<=[。！？；.!?;])\s*", text)
        chunks: list[str] = []
        buf = ""
        for s in sentences:
            if not s.strip():
                continue
            if len(buf) + len(s) <= chunk_size:
                buf += s
            else:
                if buf:
                    chunks.append(buf.strip())
                buf = s
        if buf:
            chunks.append(buf.strip())

        if len(chunks) <= 1 and len(text) > chunk_size:
            return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]
        return chunks if chunks else [text]

    @staticmethod
    def _carry_overlap(buffer: list[str], overlap_chars: int) -> list[str]:
        """计算窗口重叠部分。"""
        carry: list[str] = []
        carry_len = 0
        for seg in reversed(buffer):
            if carry_len + len(seg) <= overlap_chars:
                carry.insert(0, seg)
                carry_len += len(seg)
            else:
                break
        return carry

    @staticmethod
    def _make_chunk(
        segments: list[str],
        doc_id: str,
        idx: int,
        metadata: dict[str, Any] | None = None,
    ) -> Chunk:
        """创建 Chunk 对象，自动对齐句子边界。"""
        content = KnowledgeBase._align_chunk_end("\n\n".join(segments))
        return Chunk(
            id=KnowledgeBase._new_id(),
            document_id=doc_id,
            content=content,
            chunk_index=idx,
            metadata=metadata or {},
        )
